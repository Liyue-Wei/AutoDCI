import os
from docx import Document

def merge_word_files():
    # 输入目标文件夹路径
    folder_path = input("请输入目标文件夹路径：").replace("\\", "/").replace('\"', "").strip()
    if not os.path.isdir(folder_path):
        print("输入的路径无效，请检查后重试。")
        return

    # 输入输出文件名称
    output_file_name = input("请输入输出文件名称（不需要扩展名）：").strip()
    if not output_file_name:
        print("文件名称不能为空。")
        return
    output_file_path = os.path.join(folder_path, f"{output_file_name}.docx")

    # 创建一个新的 Word 文档
    merged_document = Document()

    # 遍历文件夹中的所有 .docx 文件
    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx"):
            file_path = os.path.join(folder_path, file_name)
            print(f"正在合并文件：{file_name}")
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                merged_document.add_paragraph(paragraph.text)
            # 添加分页符以区分不同文件内容
            merged_document.add_page_break()

    # 保存合并后的文档
    merged_document.save(output_file_path)
    print(f"所有文件已合并完成，输出文件路径：{output_file_path}")

if __name__ == "__main__":
    merge_word_files()